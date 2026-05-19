#

import pickle
from docx import Document
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from datetime import datetime, timedelta
import numpy as np
import matplotlib.pyplot as plt
import turtle
import sweetviz as sv
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

#--- Dicionario dos dados iniciais ---
artigos = {}

#--- Funcoes para gerir artigos: Registar Dados ---
def registar_dados():
    dado = input("Introduza a descrição do artigo que pretende registar: ")

    while True:
        try:
            stock = int(input("Digite a quantidade de stock: "))
            break
        except ValueError:
            print("Erro! Introduza um número inteiro válido.")
    while True:
        try:
            preco = float(input("Digite o preço do artigo: "))
            break
        except ValueError:
            print("Erro! Introduza um número inteiro válido.")
            
    fornecedor = input("Digite o fornecedor do artigo: ")
    categoria = input("Digite a categoria do artigo (ex: Padaria, Doçaria): ")
    local = input("Digite a localização no armazém (ex: A1, B2): ")

    while True:
        try:
            dia = int(input("Digite os dias até à data de validade: "))
            data_val = datetime.now() + timedelta(days=dia)
            break
        except ValueError:
            print("Erro! Introduza um número inteiro válido.")

    while True:
        try:
            min_stock = int(input("Digite o stock mínimo para alerta: "))
            break
        except ValueError:
            print("Erro! Introduza um número inteiro válido.")
    
    artigos[dado]={
        "nome": dado,
        "quantidade": stock,
        "valor": preco,
        "fornecedores": fornecedor,
        "categorias": categoria,
        "local_armz": local,
        "data_validade": data_val,
        "stock_min": min_stock
    }
    
    print(f"Dado '{dado}', Quantidade {stock} e Valor {preco} do(a) {fornecedor} registado com sucesso no {local}!")


#--- COnsultar os dados do Sistema ---
def consultar_dados():
    if not artigos:
        print("\nO sistema encontra-se sem dados")
    else:
        print("Registos do Sistema: ")
        
        for i, (dado, info) in enumerate(artigos.items(), start=1):
            print(f"SKU: {i} - Artigo: {dado} - Quantidade: {info['quantidade']} Un - Preço: {info['valor']} € - Fornecedor: {info['fornecedores']} - Categoria: {info['categorias']} - Armazém: {info['local_armz']} - Data de Validade: {info['data_validade']} - Stock Mínimo: {info['stock_min']}")

            
#--- Eliminar registos do Sistema ---
def eliminar_dados():
    if not artigos:
        print("\nA lista encontra-se vazia!")
    else:
        consultar_dados()
        try:
            indice = int(input("Indique o ID do registo a eliminar: "))
            if 1 <= indice <= len(artigos):
                #converte a chave para uma lista, apanha a chave correcta e apaga pela chave
                dado_eli = list(artigos.keys())[indice -1]
                item_eli = artigos.pop(dado_eli)
                
                print(f"\nArtigo {item_eli} eliminado com sucesso!")
            else:
                print("\nErro, o ID é invalido!")     
        except ValueError:
            print("\nErro! O ID deve ser um valor inteiro.")

          
#--- Alterar registos do Sistema ---
def alterar_dados():
    if not artigos:
        print("\nA lista encontra-se vazia! ")
    else:
        consultar_dados()
        try:
            indice = int(input("Indique o ID do registo a alterar: "))
            if 1 <= indice <= len(artigos):
                #obter a key correspondente ao artigo
                dado_novo = list(artigos.keys())[indice -1]
                print(f"Alterando o artigo {dado_novo}.")
                
                #pedir novos dados ao user
                nome_novo = input("Altere o nome: ")
                quant_novo = input("Altere a quantidade: ")
                valor_novo = input("Altere o valor: ")
                fornecedores_novo = input("Altere o fornecedor: ")
                categorias_novo = input("Altere a categoria: ")
                local_armz_novo = input("Altere a localização: ")
                data_validade_novo = input("Altere a validade: ")
                stock_min_novo = input("Altere o stock mínimo: ")
                
                #update se o user introduziu valores
                artigos[dado_novo]["nome"] = nome_novo
                if quant_novo: artigos[dado_novo]["quantidade"] = int(quant_novo)
                if valor_novo: artigos[dado_novo]["valor"] = float(valor_novo)
                if fornecedores_novo: artigos[dado_novo]["fornecedores"] = fornecedores_novo
                if categorias_novo: artigos[dado_novo]["categorias"] = categorias_novo
                if local_armz_novo: artigos[dado_novo]["local_arm"] = local_armz_novo
                if data_validade_novo: 
                    try:
                        dias = int(data_validade_novo)
                        #nova_data = datetime.now() + timedelta(days=dias)
                        artigos[dado_novo]["data_validade"] = datetime.now() + timedelta(days=dias) #nova_data.strftime("%d-%m-%Y") #%Y%m%d_%H%M%S
                    except ValueError:
                        print("\nValor inválido para validade. Deve ser um número de dias.")
                if stock_min_novo: artigos[dado_novo]["stock_min"] = int(stock_min_novo)
                
                
                print("\nO artigo foi alterado com sucesso!!")
            else:
                print("\nErro! O ID é invalido!")
        except ValueError:
            print("\nDeu erro.. o seu ID deve ser um valor inteiro..")


##--- Criar Exportações de Ficheiros
def gerar_ficheiro_pickle():
    if not artigos:
        print("\nO sistema encontra-se vazio! ")
    else:
        try:
            #parcela de codigo para apresentar a esclha da directoria para guardar um ficheiro
            file_path = input("Indique a diretoria e o nome do ficheiro a exportar: ")
            if not file_path.endswith(".pk1"):
                file_path += ".pk1"
                
            #abaixo é a criaçao do ficheiro com os dados indicados nas funçoes anteriores
            with open(file_path, "wb") as ficheiro:
                pickle.dump((artigos), ficheiro)
            print("\nFicheiro gerado com sucesso! ")
        
        except Exception as e:
            print(f"\nErro ao gerar o ficheiro: {e}")
            
            
def importar_ficheiro_pickle():
    global artigos
    
    try:
        file_path = input("Indique a diretoria e o nome do ficheiro: ")
        if not file_path.endswith(".pk1"):
            file_path += ".pk1"
            
        with open(file_path,"rb") as ficheiro:
            artigos = pickle.load(ficheiro)
        print("\nDados carregados com sucesso! ")
        
    except FileNotFoundError:
        print("\nError! Ficheiro não encontrado! ")
    except Exception as e:
        print(f"\nErro ao carregar ficheiro: {e}")
        
        
def exportar_word():
    global artigos
    
    if not artigos:
        print("\nA lista encontra-se vazia! ")
    else:
        try:
            file_path = input("Indique a diretoria e o nome do ficheiro a exportar para WORD: ")
            
            if not file_path.endswith(".docx"):
                file_path += ".docx"
                
                documento = Document()
                cabecalho = input("Coloque o cabeçalho no documento Word: ")
                documento.add_heading(cabecalho, 0)
               
                for i, (dado, info) in enumerate(artigos.items(), start=1):
                    documento.add_paragraph(f"SKU: {i} - Artigo: {dado} - Quantidade: {info['quantidade']} Un - "f"Preço: {info['valor']} € - Fornecedor: {info['fornecedores']} - "f"Categoria: {info['categorias']} - Armazém: {info['local_armz']} - "f"Data de Validade: {info['data_validade']} - Stock Mínimo: {info['stock_min']}")

                documento.save(file_path)
                print(f"\nDocumento Word '{file_path}' criado com sucesso!")
                
        except Exception as e:
            print(f"\nErro ao exportar ficheiro Word: {e}")
            
            
def exportar_pdf():
    if not artigos:
        print("\nA lista encontra-se vazia! ")
    else:
        try:
            file_path = input("Indique a diretoria e o nome do ficheiro a exportar para PDF: ")
            
            if not file_path.endswith(".pdf"):
                file_path += ".pdf"
            
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("helvetica", "B", size=14)    
            titulo = input("Indique o título do documento: ")
            pdf.cell(200, 10, text=titulo, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(5)
            
            pdf.set_font("helvetica", size=10)
            for i, (nome, info) in enumerate(artigos.items(), start=1):
                texto = (
                    f"SKU: {i} - Artigo: {nome} - Quantidade: {info['quantidade']} Un - "
                    f"Preço: {info['valor']} EUR - Fornecedor: {info['fornecedores']} - "
                    f"Categoria: {info['categorias']} - Armazém: {info['local_armz']} - "
                    f"Data de Validade: {info['data_validade']} - Stock Mínimo: {info['stock_min']}"
                )
                pdf.multi_cell(0, 10, texto)
                pdf.ln(2)   
                
            pdf.output(file_path)
            print("\nDocumento em PDF exportado com sucesso!!")
            
        except Exception as e:
            print(f"\nErro ao exportar ficheiro pdf: {e}")
           
            
def exportar_excel():
    if not artigos:
        print("\nErro: Não há dados para exportar!")
        return
    else:
        try:
            file_path = input("Indique a diretoria e o nome do ficheiro a exportar para Excel: ").strip()

            if not file_path.endswith(".xlsx"):
                file_path += ".xlsx"

            #criar o livro e a folha
            wb = Workbook()
            ws = wb.active
            ws.title = "Gestão de Artigos"

            #adicionar cabeçalhos
            ws.append(["SKU", "Artigo", "Quantidade", "Preço (€)", "Fornecedor", "Categoria", "Armazém", "Data de Validade", "Stock Mínimo"])

            #estilizar os cabeçalhos
            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center")

            #adicionar os dados
            for i, (nome, info) in enumerate(artigos.items(), start=1):
            
                #converte data para string se for datetime
                data_val = info["data_validade"]
                if hasattr(data_val, "strftime"):
                    data_val = data_val.strftime("%d-%m-%Y") #%Y%m%d_%H%M%S

                ws.append([
                    i,
                    nome,
                    info["quantidade"],
                    info["valor"],
                    info["fornecedores"],
                    info["categorias"],
                    info["local_armz"],
                    data_val,
                    info["stock_min"]
                ])

            #ajustar a largura das colunas
            for col in ws.columns:
                max_length = 0
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = (max_length + 2) * 1.2
                ws.column_dimensions[col[0].column_letter].width = adjusted_width

            #guardar o ficheiro
            wb.save(file_path)
            print(f"\nFicheiro Excel '{file_path}' criado com sucesso!")

        except Exception as e:
            print(f"\nErro ao exportar ficheiro Excel: {e}")
            
                  
###--- Stats ---
def calcular_valor_total_stock():
    if not artigos:
        print("\nNão existem dados.")
        return
    
    valor_total = 0
    
    for nome, info in artigos.items():
        valor_total += info["quantidade"] * info["valor"]
    
    print(f"\nO valor total do stock é: {valor_total:.2f}€")
    

def verificar_stock_min():
    if not artigos:
        print("\nNão existem dados.")
        return
    
    alerta = False
            
    for dado, info in artigos.items():
        #2print(dado, info, type(info))
        
        if info['quantidade'] <= info['stock_min']:
            print(f"ALERTA: {dado} (Stock: {info['quantidade']}, Mínimo: {info['stock_min']})")
            alerta = True
    if not alerta:
        print("\nNão existem artigos no stock mínimo.")

  
def estatisticas_precos():
    if not artigos:
        print("\nNão há dados de preços!")
        return
    
    preco_array = np.array([info['valor'] for info in artigos.values()])
    
    media = np.mean(preco_array)
    mediana = np.median(preco_array)
    desvio_padrao = np.std(preco_array)
    min_preco = np.min(preco_array)
    max_preco = np.max(preco_array)

    print("\n--- Estatísticas de Preços ---")
    print(f"Média: {round(media, 3)} €")
    print(f"Mediana: {round(mediana, 2)}€")
    print(f"Desvio Padrão: {desvio_padrao:.2f}€")
    print(f"Preço Mínimo: {min_preco:.2f}€")
    print(f"Preço Máximo: {max_preco:.2f}€")
    
    print("\n--- Valor total por artigo ---")
    for nome, info in artigos.items():
        valor_total = info["quantidade"] * info["valor"]
        print(f"{nome}: {valor_total:.2f} €  (Stock: {info['quantidade']}, Preço: {info['valor']} €)")
 

def gerar_relatorio_sweetviz():
    if not artigos:
        print("\nErro: Não há dados para gerar o relatório!")
        return
    
    #Criar um DataFrame com os dados
    tabela = {
        "Artigo": [],
        "Quantidade": [],
        "Preço (€)": [],
        "Fornecedor": [],
        "Categoria": [],
        "Localização": [],
        "Data de Validade": [],
        "Stock Mínimo": [] 
    }
    
    for nome, info in artigos.items():
        tabela["Artigo"].append(nome)
        tabela["Quantidade"].append(info.get("quantidade"))
        tabela["Preço (€)"].append(info.get("valor"))
        tabela["Fornecedor"].append(info.get("fornecedores"))
        tabela["Categoria"].append(info.get("categorias"))
        tabela["Localização"].append(info.get("local_armz"))
        tabela["Data de Validade"].append(info.get("data_validade"))
        tabela["Stock Mínimo"].append(info.get("min_stock"))

    df = pd.DataFrame(tabela)

    #gerar relatorio
    relatorio = sv.analyze(df)

    #guardar em HTML
    timestamp = datetime.now().strftime("%d-%m-%Y")
    nome_ficheiro = f"relatorio_sweetviz_{timestamp}.html"
    relatorio.show_html(nome_ficheiro)

    print(f"\nRelatório Sweetviz gerado com sucesso: '{nome_ficheiro}'")
    print(f"Podes abrir o ficheiro no teu browser para ver o relatório.")  

 
##--- Gráficos ---
def gerar_grafico_circular():
    if not artigos:
        print("\nNão existem dados para criar gráficos.")
        return
    
    nomes = list(artigos.keys())
    quantidades = [info["quantidade"] for info in artigos.values()]
    
    plt.pie(quantidades, labels=nomes, autopct='%1.1f%%')
    plt.title("Distribuição das Quantidades por Artigo", fontsize=16)
    plt.show()
    
def gerar_grafico_barras():
    if not artigos:
        print("\nNão existem dados para criar gráficos.")
        return
    
    #dict para adicionar valores
    valores_categorias = {}
    for info in artigos.values():
        categoria=info["categorias"]
        valor=info["valor"]*info["quantidade"]
        
        if categoria in valores_categorias:
            valores_categorias[categoria]+=valor
        else:
            valores_categorias[categoria]=valor
            
    categorias = list(valores_categorias.keys())
    valores = list(valores_categorias.values())
    
    plt.bar(categorias, valores)
    plt.xlabel("Eixo X - Categoria")
    plt.ylabel("Eixo Y - Preço €")
    plt.title("Distribuição Valorizada por Categoria ")
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.show()
    
def desenhar_diagrama_valor():
    if not artigos:
        print("\nNão existem dados para desenhar.")
        return
    else:
        
        nomes = list(artigos.keys())
        valores = [info["valor"] for info in artigos.values()]
        
        screen = turtle.Screen()
        screen.title("Desenhar Preços dos Artigos")
        screen.setup(width=800, height=600)
        screen.bgcolor("azure")
        
        t = turtle.Turtle()
        t.speed(3.5)
        #t.hideturtle() #esconde o cursos turtle?
        t.width(3)
        
        #cores das barras
        cores = ["#FF6B6B", "#4ECDC4", "#45B7D1", "#96CEB4", "#FFEAA7", "#DDA0DD", "#98D8C8"]
        #logica para o valor máximo para escalar as barras
        max_valor = max(valores)
        escala = 600 / max_valor  # Escala para caber na janela (600 pixels de largura)
        #titulo
        t.penup()
        t.goto(0, 280)
        t.write("Preços dos Artigos (€)", align="center", font=("Helvetica", 14, "bold"))
        t.color("black")
        #linhas verticais
        t.penup()
        t.goto(-300, -75)
        t.pendown()
        t.color("gray")
        for i in range(0, int(max_valor) + 1, max(1, int(max_valor) // 10)):
            t.penup()
            t.goto(-300 + i * escala, 250)
            t.pendown()
            t.goto(-300 + i * escala, -100)
        t.penup()
        
        #eixo Y (artigos)
        t.penup()
        for i, dado in enumerate(nomes):
            t.goto(-300, 250 - i * 50)
            t.write(dado, align="right", font=("Helvetica", 12, "normal"))
            
        #eixo X (valor)
        t.penup()
        t.goto(-300, -100)
        t.pendown()
        t.goto(300, -100)
        for i in range(0, int(max_valor) + 1, max(1, int(max_valor) // 10)):
            t.penup()
            t.goto(-300 + i * escala, -100)
            t.pendown()
            t.goto(-300 + i * escala, -120)
            t.penup()
            t.write(f"{i:.2f}€", align="center", font=("Helvetica", 10, "normal"))

        #desenhar as barras
        for i, preco in enumerate(valores):
            t.penup()
            t.goto(-300, 275 - i * 50)
            t.pendown()
            t.color(cores[i % len(cores)])
            t.begin_fill()
            t.forward(preco * escala)
            t.right(90)
            t.forward(40)
            t.right(90)
            t.forward(preco * escala)
            t.right(90)
            t.forward(40)
            t.right(90)
            t.end_fill()
            
            t.penup()
            t.goto(-300 + preco * escala + 10, 200 - i * 50 + 10)
            t.write(f"{preco:.2f}€", align="left", font=("Helvetica", 10, "normal"))
        #turtle.done()
        
#df = pd.DataFrame.from_dict(artigos, orient='index')
#df.index.name = 'Artigo'
#df.reset_index(inplace=True)         

def filtrar_por_categoria():
    if not artigos:
        print("\nNão há dados para filtrar!")
        return

    df = pd.DataFrame.from_dict(artigos, orient='index')
    df.index.name = 'Artigo'
    df.reset_index(inplace=True)

    categoria = input("\nDigite a categoria a filtrar: ")
    df_filtrado = df[df['categorias'].str.lower() == categoria.lower()]

    if df_filtrado.empty:
        print(f"\nNão foram encontrados artigos na categoria '{categoria}'.")
    else:
        print(f"\n--- Artigos na categoria '{categoria}' ---")
        print(df_filtrado.to_string(index=False))
        
def filtrar_por_fornecedor():
    if not artigos:
        print("\nNão há dados para filtrar!")
        return

    df = pd.DataFrame.from_dict(artigos, orient='index')
    df.index.name = 'Artigo'
    df.reset_index(inplace=True)

    fornecedor = input("\nDigite o fornecedor a filtrar: ")
    df_filtrado = df[df['fornecedores'].str.lower() == fornecedor.lower()]

    if df_filtrado.empty:
        print(f"\nNão foram encontrados artigos do fornecedor '{fornecedor}'.")
    else:
        print(f"\n--- Artigos do fornecedor '{fornecedor}' ---")
        print(df_filtrado.to_string(index=False))
        
def ordenar_por_preco():
    if not artigos:
        print("\nNão há dados para ordenar!")
        return

    df = pd.DataFrame.from_dict(artigos, orient='index')
    df.index.name = 'Artigo'
    df.reset_index(inplace=True)

    ordem = input("\nOrdenar por preço (1 - Crescente, 2 - Decrescente): ")
    if ordem == "1":
        df_ordenado = df.sort_values('valor', ascending=True)
    elif ordem == "2":
        df_ordenado = df.sort_values('valor', ascending=False)
    else:
        print("\nOpção inválida!")
        return

    print("\n--- Artigos Ordenados por Preço ---")
    print(df_ordenado.to_string(index=False))       
        
def adicionar_valor_total():
    if not artigos:
        print("\nNão há dados!")
        return

    df = pd.DataFrame.from_dict(artigos, orient='index')
    df.index.name = 'Artigo'
    df.reset_index(inplace=True)

    # Adicionar coluna de valor total
    df['valor_total'] = df['quantidade'] * df['valor']

    print("\n--- Artigos com Valor Total ---")
    print(df.to_string(index=False))

    return df

def analise_stock_localizacao():
    if not artigos:
        print("\nNão há dados para análise!")
        return

    df = pd.DataFrame.from_dict(artigos, orient='index')
    df.index.name = 'Artigo'
    df.reset_index(inplace=True)

    # Agrupar por localização
    stock_por_local = df.groupby('local_armz').agg({
        'quantidade': 'sum',
        'valor': 'mean',
        'Artigo': lambda x: ', '.join(x)  # Lista de artigos por localização
    }).reset_index()

    stock_por_local.rename(columns={
        'quantidade': 'Quantidade Total',
        'valor': 'Preço Médio (€)',
        'Artigo': 'Artigos'
    }, inplace=True)

    print("\n--- Stock por Localização ---")
    print(stock_por_local.to_string(index=False))

    # Gráfico de barras
    plt.figure(figsize=(10, 6))
    plt.bar(stock_por_local['local_armz'], stock_por_local['Quantidade Total'], color='skyblue')
    plt.xlabel("Localização no Armazém")
    plt.ylabel("Quantidade Total")
    plt.title("Stock por Localização no Armazém")
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.show()

def analise_validade():
    if not artigos:
        print("\nNão há dados para análise!")
        return

    df = pd.DataFrame.from_dict(artigos, orient='index')
    df.index.name = 'Artigo'
    df.reset_index(inplace=True)

    # Calcular dias até à validade
    df['dias_ate_validade'] = (df['data_validade'] - datetime.now()).dt.days

    # Artigos que expiram em menos de 7 dias
    expiracao_proxima = df[df['dias_ate_validade'] <= 7]

    print("\n--- Análise de Validade ---")
    print("\nDias até à validade por artigo:")
    print(df[['Artigo', 'dias_ate_validade', 'data_validade']].to_string(index=False))

    if not expiracao_proxima.empty:
        print("\nATENÇÃO: Os seguintes artigos expiram em menos de 7 dias:")
        print(expiracao_proxima[['Artigo', 'dias_ate_validade', 'data_validade']].to_string(index=False))
    else:
        print("\nNão há artigos a expirar em menos de 7 dias.")

    # Gráfico de barras
    plt.figure(figsize=(10, 6))
    plt.bar(df['Artigo'], df['dias_ate_validade'], color='lightcoral')
    plt.xlabel("Artigo")
    plt.ylabel("Dias até à Validade")
    plt.title("Dias até à Validade por Artigo")
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.show()                

def exportar_dataframe_excel(df, nome_ficheiro=None):
    if df.empty:
        print("\nO DataFrame está vazio!")
        return

    if nome_ficheiro is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nome_ficheiro = f"dataframe_{timestamp}.xlsx"

    try:
        df.to_excel(nome_ficheiro, index=False)
        print(f"\nDataFrame exportado para '{nome_ficheiro}' com sucesso!")
    except Exception as e:
        print(f"\nErro ao exportar DataFrame: {e}")
        
def pesquisa_avancada():
    if not artigos:
        print("\nNão há dados para pesquisar!")
        return

    df = pd.DataFrame.from_dict(artigos, orient='index')
    df.index.name = 'Artigo'
    df.reset_index(inplace=True)

    print("\n--- Pesquisa Avançada ---")
    print("1. Filtrar por categoria")
    print("2. Filtrar por fornecedor")
    print("3. Filtrar por localização")
    print("4. Filtrar por stock mínimo")
    print("5. Filtrar por data de validade (dias)")
    opcao = input("\nSelecione uma opção: ")

    if opcao == "1":
        categoria = input("Digite a categoria: ")
        df_filtrado = df[df['categorias'].str.lower() == categoria.lower()]
    elif opcao == "2":
        fornecedor = input("Digite o fornecedor: ")
        df_filtrado = df[df['fornecedores'].str.lower() == fornecedor.lower()]
    elif opcao == "3":
        local = input("Digite a localização: ")
        df_filtrado = df[df['local_armz'].str.lower() == local.lower()]
    elif opcao == "4":
        stock_min = int(input("Digite o stock mínimo: "))
        df_filtrado = df[df['quantidade'] <= stock_min]
    elif opcao == "5":
        dias = int(input("Digite os dias até à validade: "))
        df['dias_ate_validade'] = (df['data_validade'] - datetime.now()).dt.days
        df_filtrado = df[df['dias_ate_validade'] <= dias]
    else:
        print("\nOpção inválida!")
        return

    if df_filtrado.empty:
        print("\nNão foram encontrados resultados.")
    else:
        print("\n--- Resultados da Pesquisa ---")
        print(df_filtrado.to_string(index=False))

        # Opção para exportar os resultados
        exportar = input("\nDeseja exportar os resultados para Excel? (s/n): ").lower()
        if exportar == 's':
            exportar_dataframe_excel(df_filtrado)
            
                    
        

##--- Sub Menus ---
def submenu_gestao():
    while True:
        print("=" * 50)
        print("  SUBMENU DE GESTÃO DE DADOS  ")
        print("=" * 50)
        print(" 1. Registar artigos")
        print(" 2. Consultar artigos")
        print(" 3. Eliminar artigos")
        print(" 4. Alterar artigos")
        print(" 0. Voltar ao menu principal")
        print("=" * 50)
        escolha = input("Indique a opção a executar: ")

        if escolha == "1":
            registar_dados()
        elif escolha == "2":
            consultar_dados()
        elif escolha == "3":
            eliminar_dados()
        elif escolha == "4":
            alterar_dados()
        elif escolha == "0":
            print("Programa Encerrado.")
            break
        else:
            print("Erro! Opção Inválida! Só pode digitar as opções apresentadas.")
        
def submenu_exportacao():
    while True:
        print("=" * 50)
        print("  SUBMENU DE IMPORTAÇÃO-EXPORTAÇÃO  ")
        print("=" * 50)
        print(" 1. Gerar ficheiro (Pickle)")
        print(" 2. Importar ficheiro (Pickle)")
        print(" 3. Exportar em Word")
        print(" 4. Exportar em PDF")
        print(" 5. Exportar em Excel")
        print(" 0. Voltar ao menu principal")
        print("=" * 50)
        escolha = input("Indique a opção a executar: ")

        if escolha == "1":
            gerar_ficheiro_pickle()
        elif escolha == "2":
            importar_ficheiro_pickle()
        elif escolha == "3":
            exportar_word()
        elif escolha == "4":
            exportar_pdf()
        elif escolha == "5":
            exportar_excel()
        elif escolha == "0":
            print("Programa Encerrado.")
            break
        else:
            print("Erro! Opção Inválida! Só pode digitar as opções apresentadas.")
  
def submenu_stocks():
    while True:
        print("=" * 50)
        print("  SUBMENU DE STOCKS-GRÁFICOS-RELATÓRIOS  ")
        print("=" * 50)
        print(" 1. Cacular valor total de stocks")
        print(" 2. Verificar Produtos em Stock Minimo")
        print(" 3. Estatísticas de Preços")
        print(" 4. Obter gráfico circular de Quantidades por Produto")
        print(" 5. Obter gráfico barras de Valor por Categorias")
        print(" 6. Desenhar Diagrama Turtle")
        print(" 7. Gerar relatório Sweetviz")
        print(" 0. Voltar ao menu principal")
        print("=" * 50)
        escolha = input("Indique a opção a executar: ")

        if escolha == "1":
            calcular_valor_total_stock()
        elif escolha == "2":
            verificar_stock_min()
        elif escolha == "3":
            estatisticas_precos()
        elif escolha == "4":
            gerar_grafico_circular()
        elif escolha == "5":
            gerar_grafico_barras()
        elif escolha == "6":
            desenhar_diagrama_valor()
        elif escolha == "7":
            gerar_relatorio_sweetviz()
        elif escolha == "0":
            print("Programa Encerrado.")
            break
        else:
            print("Erro! Opção Inválida! Só pode digitar as opções apresentadas.")


def submenu_analise_dados():
    while True:
        print("=" * 50)
        print("  SUBMENU DE ANÁLISE DE DADOS  ")
        print("=" * 50)
        print(" 1. Filtrar por categoria")
        print(" 2. Filtrar por fornecedor")
        print(" 3. Ordenar por preço")
        print(" 4. Análise de stock por localização")
        print(" 5. Análise de validade")
        print(" 6. Pesquisa avançada")
        print(" 7. Ver DataFrame completo")
        print(" 0. Voltar ao menu principal")
        print("=" * 50)
        escolha = input("Indique a opção a executar: ")

        if escolha == "1":
            filtrar_por_categoria()
        elif escolha == "2":
            filtrar_por_fornecedor()
        elif escolha == "3":
            ordenar_por_preco()
        elif escolha == "4":
            analise_stock_localizacao()
        elif escolha == "5":
            analise_validade()
        elif escolha == "6":
            pesquisa_avancada()
        elif escolha == "7":
            df = pd.DataFrame.from_dict(artigos, orient='index')
            df.index.name = 'Artigo'
            df.reset_index(inplace=True)
            print("\n--- DataFrame Completo ---")
            print(df.to_string(index=False))
        elif escolha == "0":
            break
        else:
            print("\nOpção inválida!")




    
##--- Menu ---
def menu_programa():
    while True:
        print("=" * 50)
        print("  BEM-VINDO(A) AO PROGRAMA DE GESTÃO DE ARTIGOS  ")
        print("=" * 50)
        print("Este programa permite o registo de artigos, quantidade e preço.\n")
        print(" ===== MENU PRINCIPAL =====")
        print(" 1. Gestão de Dados")
        print(" 2. Importação e Exportação de Dados")
        print(" 3. Gestão de Stocks, Dados Estatísticos e Gráficos")
        print(" 4. Gestão Dinâmica e Análise dos Dados")
        print(" 0. Terminar programa")
        print("=" * 50)
        escolha = input("Indique a opção a executar: ")

        if escolha == "1":
            submenu_gestao()
        elif escolha == "2":
            submenu_exportacao()
        elif escolha == "3":
            submenu_stocks()
        elif escolha == "4":
            submenu_analise_dados()   
        elif escolha == "0":
            print("Programa Encerrado.")
            break
        else:
            print("Erro! Opção Inválida! Só pode digitar as opções apresentadas.")
            
if __name__ == "__main__":
    menu_programa()